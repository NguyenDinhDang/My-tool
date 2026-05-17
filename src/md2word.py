import argparse
import os
import sys
import re
import base64
import urllib.request
import urllib.error
import socket
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import docx.opc.constants

for stream in (sys.stdout, sys.stderr):
    if hasattr(stream, 'reconfigure'):
        stream.reconfigure(encoding='utf-8', errors='replace')

_MERMAID_INK_DISABLED = False

def add_hyperlink(paragraph, text, url):
    """Hàm hỗ trợ chèn Hyperlink thật vào Word (có thể click được)"""
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Định dạng link màu xanh, gạch chân
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def parse_inline_text(paragraph, text):
    """Xử lý inline formatting: Bold, Inline Code, và Links"""
    # Pattern tìm links: [text](url)
    link_pattern = r'\[([\]]+)\]\(([)]+)\)'
    
    # Tách đoạn text thành các phần có chứa link và không chứa link
    parts_by_link = re.split(link_pattern, text)
    
    i = 0
    while i < len(parts_by_link):
        if i % 3 == 0:  # Text bình thường (có thể chứa bold/code)
            sub_text = parts_by_link[i]
            # Tách tiếp Bold và Inline code
            parts = re.split(r'(\*\*.*?\*\*|`.*?`)', sub_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = paragraph.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('`') and part.endswith('`'):
                    run = paragraph.add_run(part[1:-1])
                    run.font.name = 'Courier New'
                    run.font.color.rgb = RGBColor(199, 37, 78)
                    run.font.size = Pt(10)
                else:
                    paragraph.add_run(part)
            i += 1
        else: # Cụm link [text](url)
            link_text = parts_by_link[i]
            link_url = parts_by_link[i+1]
            add_hyperlink(paragraph, link_text, link_url)
            i += 2

def setup_styles(doc):
    """Thiết lập màu sắc và font chữ mặc định"""
    styles = doc.styles
    
    h1 = styles['Heading 1']
    h1.font.name = 'Arial'
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 51, 102)

    h2 = styles['Heading 2']
    h2.font.name = 'Arial'
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 102, 204)

    h3 = styles['Heading 3']
    h3.font.name = 'Arial'
    h3.font.size = Pt(12)
    h3.font.bold = True

    normal = styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)

def build_table(doc, table_lines):
    """Chuyển đổi các dòng text Markdown thành Table trong Word"""
    if not table_lines or len(table_lines) < 2: return
    
    # Lấy số cột từ dòng đầu tiên
    headers = [col.strip() for col in table_lines[0].split('|') if col.strip()]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Fill Header
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        
    # Bỏ qua dòng phân cách (vd: |---|---|)
    # Fill Data
    for row_line in table_lines[2:]:
        cols = [col.strip() for col in row_line.split('|')[1:-1]]
        if cols:
            row_cells = table.add_row().cells
            for i, col in enumerate(cols):
                if i < len(row_cells):
                    parse_inline_text(row_cells[i].paragraphs[0], col)

def _render_mermaid_diagram_legacy(doc, mermaid_code):
    """
    Render Mermaid diagram thành hình ảnh và nhúng vào Word
    
    Cách 1: Dùng mermaid-cli (nếu cài Node.js)
    Cách 2: Nhúng code dưới dạng text với link mermaid.live
    """
    import subprocess
    import tempfile
    
    try:
        # Cách 1: Thử dùng mmdc (mermaid-cli)
        print(f" Đang render Mermaid diagram (mmdc)...")
        
        with tempfile.TemporaryDirectory() as temp_dir:
            # Lưu mermaid code vào temp file
            temp_md = os.path.join(temp_dir, 'diagram.mmd')
            temp_png = os.path.join(temp_dir, 'diagram.png')
            
            with open(temp_md, 'w', encoding='utf-8') as f:
                f.write(mermaid_code)
            
            # Cách chạy mmdc qua shell (để tìm được từ PATH)
            try:
                # Thử cách 1: Chạy qua shell (Windows)
                result = subprocess.run(
                    f'mmdc -i "{temp_md}" -o "{temp_png}" -w 800',
                    shell=True,
                    capture_output=True,
                    timeout=30,
                    text=True
                )
                
                if result.returncode == 0 and os.path.exists(temp_png):
                    # Nhúng hình vào Word
                    p = doc.add_paragraph()
                    p_format = p.paragraph_format
                    p_format.alignment = 1  # Center
                    p_format.space_before = Pt(6)
                    p_format.space_after = Pt(6)
                    
                    p.add_run().add_picture(temp_png, width=Inches(5.5))
                    print(f"   Mermaid diagram đã được render thành hình!")
                    return True
                else:
                    # Nếu command thất bại, hiển thị lỗi
                    if result.stderr:
                        print(f"   mmdc lỗi: {result.stderr[:100]}")
                    raise Exception("mmdc command failed")
                    
            except (FileNotFoundError, subprocess.TimeoutExpired) as e:
                print(f"  mmdc không tìm thấy hoặc timeout")
                raise
                
    except Exception as e:
        # Fallback: Nhúng mermaid code dưới dạng code block
        print(f"   Nhúng code Mermaid (copy vào mermaid.live để xem)")
        
        p = doc.add_paragraph()
        p_format = p.paragraph_format
        p_format.left_indent = Inches(0.3)
        p_format.right_indent = Inches(0.3)
        p_format.space_before = Pt(6)
        p_format.space_after = Pt(6)
        
        # Thêm ghi chú
        info_run = p.add_run("[ Mermaid Diagram - Dán vào mermaid.live để xem]\n\n")
        info_run.italic = True
        info_run.font.color.rgb = RGBColor(100, 100, 100)
        
        # Thêm code
        code_run = p.add_run(mermaid_code)
        code_run.font.name = 'Courier New'
        code_run.font.size = Pt(8)
        code_run.font.color.rgb = RGBColor(51, 51, 51)
        
        # Background color
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'FFF3CD')
        p._element.get_or_add_pPr().append(shading)
        
        # Thêm link đến mermaid.live
        p2 = doc.add_paragraph()
        p2.add_run(" Xem trực tiếp: ")
        add_hyperlink(p2, "mermaid.live", "https://mermaid.live")
        
        print(f"     Code Mermaid đã được nhúng")
        return False

def _add_mermaid_picture(doc, image_path, width=Inches(5.5)):
    p = doc.add_paragraph()
    p_format = p.paragraph_format
    p_format.alignment = 1  # Center
    p_format.space_before = Pt(6)
    p_format.space_after = Pt(6)
    p.add_run().add_picture(image_path, width=width)

def _render_mermaid_with_mmdc(mermaid_code, output_png):
    import shutil
    import subprocess
    import tempfile

    mmdc_path = shutil.which('mmdc') or shutil.which('mmdc.cmd')
    if not mmdc_path:
        raise FileNotFoundError('mmdc not found')

    with tempfile.NamedTemporaryFile('w', suffix='.mmd', encoding='utf-8', delete=False) as temp_md:
        temp_md.write(mermaid_code)
        temp_md_path = temp_md.name

    try:
        result = subprocess.run(
            [mmdc_path, '-i', temp_md_path, '-o', output_png, '-w', '1200', '-b', 'transparent'],
            capture_output=True,
            timeout=45,
            text=True
        )
        if result.returncode != 0 or not os.path.exists(output_png):
            error = result.stderr.strip() or result.stdout.strip() or 'mmdc command failed'
            raise RuntimeError(error[:300])
    finally:
        try:
            os.remove(temp_md_path)
        except OSError:
            pass

def _render_mermaid_with_ink(mermaid_code, output_png):
    if _MERMAID_INK_DISABLED:
        raise RuntimeError('Mermaid Ink was disabled after an earlier network failure')

    encoded = base64.urlsafe_b64encode(mermaid_code.encode('utf-8')).decode('ascii')
    url = f'https://mermaid.ink/img/{encoded}?type=png&bgColor=transparent'
    request = urllib.request.Request(url, headers={'User-Agent': 'SecurityAutomationToolkit-md2word/1.0'})

    with urllib.request.urlopen(request, timeout=8) as response:
        image_data = response.read()

    if not image_data.startswith(b'\x89PNG'):
        raise RuntimeError('Mermaid Ink did not return a PNG image')

    with open(output_png, 'wb') as image_file:
        image_file.write(image_data)

def _add_mermaid_code_fallback(doc, mermaid_code):
    p = doc.add_paragraph()
    p_format = p.paragraph_format
    p_format.left_indent = Inches(0.3)
    p_format.right_indent = Inches(0.3)
    p_format.space_before = Pt(6)
    p_format.space_after = Pt(6)

    info_run = p.add_run("[Mermaid Diagram - paste into mermaid.live to view]\n\n")
    info_run.italic = True
    info_run.font.color.rgb = RGBColor(100, 100, 100)

    code_run = p.add_run(mermaid_code)
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(8)
    code_run.font.color.rgb = RGBColor(51, 51, 51)

    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'FFF3CD')
    p._element.get_or_add_pPr().append(shading)

    p2 = doc.add_paragraph()
    p2.add_run("View online: ")
    add_hyperlink(p2, "mermaid.live", "https://mermaid.live")

def render_mermaid_diagram(doc, mermaid_code):
    """
    Render Mermaid to a PNG image in Word.

    Tries local mermaid-cli first, then Mermaid Ink online API, then falls
    back to embedding the Mermaid source code with a mermaid.live link.
    """
    import tempfile
    global _MERMAID_INK_DISABLED

    if not mermaid_code.strip():
        return False

    with tempfile.TemporaryDirectory() as temp_dir:
        temp_png = os.path.join(temp_dir, 'diagram.png')

        try:
            print("  Rendering Mermaid diagram with mmdc...")
            _render_mermaid_with_mmdc(mermaid_code, temp_png)
            _add_mermaid_picture(doc, temp_png)
            print("    Mermaid diagram rendered as image.")
            return True
        except Exception as mmdc_error:
            print(f"    mmdc unavailable or failed: {mmdc_error}")

        try:
            print("  Rendering Mermaid diagram with mermaid.ink...")
            _render_mermaid_with_ink(mermaid_code, temp_png)
            _add_mermaid_picture(doc, temp_png)
            print("    Mermaid diagram rendered as image via mermaid.ink.")
            return True
        except (urllib.error.URLError, TimeoutError, socket.timeout, OSError, RuntimeError) as ink_error:
            print(f"    mermaid.ink failed: {ink_error}")
            if isinstance(ink_error, (urllib.error.URLError, TimeoutError, socket.timeout, OSError)):
                _MERMAID_INK_DISABLED = True

    print("  Embedding Mermaid code fallback.")
    _add_mermaid_code_fallback(doc, mermaid_code)
    return False

def convert_md_to_docx(input_file, output_file=None):
    if not os.path.exists(input_file):
        print(f"Lỗi: Không tìm thấy file '{input_file}'")
        sys.exit(1)

    if not output_file:
        output_file = os.path.splitext(input_file)[0] + ".docx"

    with open(input_file, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()
    setup_styles(doc)

    lines = content.split('\n')
    in_code_block = False
    code_lines = []
    code_type = 'text'  # 'mermaid' hoặc 'text'
    table_lines = []

    print(f"Đang xử lý: {input_file} -> {output_file}")

    for line in lines:
        # 1. Xử lý Code Block
        if line.startswith('```'):
            if in_code_block:
                # End of code block
                if code_type == 'mermaid':
                    # Render Mermaid diagram
                    mermaid_code = '\n'.join(code_lines)
                    render_mermaid_diagram(doc, mermaid_code)
                else:
                    # Regular code block
                    if code_lines:
                        code_text = '\n'.join(code_lines)
                        p = doc.add_paragraph()
                        p_format = p.paragraph_format
                        p_format.left_indent = Inches(0.3)
                        p_format.right_indent = Inches(0.3)
                        p_format.space_before = Pt(6)
                        p_format.space_after = Pt(6)
                        
                        run = p.add_run(code_text)
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(51, 51, 51)
                        
                        # Thêm Shading nền xám
                        shading = OxmlElement('w:shd')
                        shading.set(qn('w:fill'), 'F5F5F5')
                        p._element.get_or_add_pPr().append(shading)
                
                code_lines = []
                code_type = 'text'
                in_code_block = False
            else:
                # Start of code block
                in_code_block = True
                # Check if it's mermaid block: ```mermaid
                lang = line[3:].strip()
                code_type = 'mermaid' if lang.lower() == 'mermaid' else 'text'
            continue
            
        if in_code_block:
            code_lines.append(line)
            continue

        # 2. Xử lý Bảng (Table)
        if line.strip().startswith('|') and line.strip().endswith('|'):
            table_lines.append(line.strip())
            continue
        elif table_lines:
            build_table(doc, table_lines)
            table_lines = []

        # 3. Xử lý Headings
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
            
        # 4. Xử lý Blockquotes
        elif line.startswith('> '):
            p = doc.add_paragraph(style='Quote')
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(line[2:])
            run.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)

        # 5. Xử lý Bullet Points
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            indent_level = (len(line) - len(line.lstrip())) // 2
            text = line.strip()[2:]
            p = doc.add_paragraph(style='List Bullet')
            if indent_level > 0:
                p.paragraph_format.left_indent = Inches(0.25 * (indent_level + 1))
            parse_inline_text(p, text)

        # 6. Xử lý Numbered Lists
        elif re.match(r'\d+\.\s', line.strip()):
            text = re.sub(r'\d+\.\s', '', line.strip())
            p = doc.add_paragraph(style='List Number')
            parse_inline_text(p, text)

        # 7. Xử lý Dòng trống
        elif line.strip() == '':
            if doc.paragraphs and doc.paragraphs[-1].text.strip() != '':
                doc.add_paragraph()

        # 8. Xử lý Text thường
        else:
            if line.strip():
                p = doc.add_paragraph()
                parse_inline_text(p, line)

    # Đảm bảo in bảng ra nếu file kết thúc ngay sau khi hết bảng
    if table_lines:
        build_table(doc, table_lines)

    try:
        doc.save(output_file)
        print(f"Hoàn tất! Đã lưu tại: {output_file}")
    except Exception as e:
        print(f"Lỗi khi lưu file: {e}")
        print("Gợi ý: Hãy chắc chắn file Word này đang không được mở bởi phần mềm khác.")

def main():
    parser = argparse.ArgumentParser(description="Tool chuyển MD sang Word (Hỗ trợ Table, Code block, Links)")
    parser.add_argument("input", help="Đường dẫn file Markdown (.md)")
    parser.add_argument("-o", "--output", help="Đường dẫn file Word (.docx) đầu ra", default=None)
    args = parser.parse_args()
    convert_md_to_docx(args.input, args.output)

if __name__ == "__main__":
    main()
